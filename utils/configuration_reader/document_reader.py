from typing import List
from pathlib import Path
from data_classes.base_models.source_types import DocxDocument, TableElement, DocElement
import docx


    

def read_word_file(doc_path:Path) -> DocxDocument:

    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        raise ValueError(f"Could not open file {doc_path}. Error {e}")
    

    doc_data = DocxDocument(file_name=doc_path.name)

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            doc_data.paragraphs.append(paragraph.text.strip())
            
    for table in doc.tables:
        current_table = []
        for row in table.rows:
            current_row = []

            for cell in row.cells:
                clean_text = cell.text.strip().replace("\n", " ")
                current_row.append(clean_text) 
            
            current_table.append(current_row)

        doc_data.tables.append(current_table)

    return doc_data