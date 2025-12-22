from typing import Union
from pathlib import Path

import docx

from data.base_models.docx_document import DocxDocument, DocElement, TableElement
from .base_reader import BaseFileReader


def read_word_file(doc_path: Path) -> DocxDocument:
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        raise ValueError(f"Could not open file {doc_path}. Error {e}")
    
    doc_data = DocxDocument(source_name=doc_path.name, 
                            source_path=doc_path, 
                            content=[])
    
    # Process paragraphs and tables in document order
    for element in doc.element.body:
        # Handle paragraphs
        if element.tag.endswith('p'):
            # Find the corresponding paragraph object
            for para in doc.paragraphs:
                if para._element == element:
                    text = para.text.strip()
                    if text:
                        # Determine if it's a heading
                        if para.style.name.startswith('Heading'):
                            level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                            doc_element = DocElement(
                                type="heading",
                                text=text,
                                metadata={"level": level, "style": para.style.name}
                            )
                        else:
                            # Regular paragraph
                            metadata = {}
                            # Check for basic formatting
                            if para.runs:
                                metadata["bold"] = any(run.bold for run in para.runs)
                                metadata["italic"] = any(run.italic for run in para.runs)
                            
                            doc_element = DocElement(
                                type="paragraph",
                                text=text,
                                metadata=metadata
                            )
                        doc_data.content.append(doc_element)
                    break
        
        # Handle tables
        elif element.tag.endswith('tbl'):
            # Find the corresponding table object
            for table in doc.tables:
                if table._element == element:
                    table_data = []
                    for row in table.rows:
                        current_row = []
                        for cell in row.cells:
                            clean_text = cell.text.strip().replace("\n", " ")
                            current_row.append(clean_text)
                        table_data.append(current_row)
                    
                    # Create concatenated text representation for the table
                    table_text = "\n".join([" | ".join(row) for row in table_data])
                    
                    table_element = TableElement(
                        type="table",
                        text=table_text,
                        data=table_data,
                        metadata={"rows": len(table_data), "cols": len(table_data[0]) if table_data else 0}
                    )
                    doc_data.content.append(table_element)
                    break
    
    return doc_data


class DocReader(BaseFileReader):
    def read(self, doc_path:Union[Path, str]) -> DocxDocument:
        doc_path = Path(doc_path)
        return read_word_file(doc_path)
